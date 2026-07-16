import docx
from docx.shared import Inches, Pt
import re

def parse_markdown_to_docx(md_path, docx_path):
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_headers = []
    table_rows = []
    
    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_headers = [cell.strip() for cell in stripped.split('|')[1:-1]]
                table_rows = []
            else:
                if re.match(r'^\|[\s\:\-\|]+\|$', stripped):
                    continue
                row_cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
                table_rows.append(row_cells)
            continue
        else:
            if in_table:
                in_table = False
                if table_headers and table_rows:
                    table = doc.add_table(rows=1, cols=len(table_headers))
                    table.style = 'Light Shading Accent 1'
                    hdr_cells = table.rows[0].cells
                    for i, header in enumerate(table_headers):
                        hdr_cells[i].text = header
                    for row_data in table_rows:
                        row_cells = table.add_row().cells
                        for i, cell_val in enumerate(row_data):
                            if i < len(row_cells):
                                clean_val = re.sub(r'\*\*(.*?)\*\*|_(.*?)_', r'\1\2', cell_val)
                                row_cells[i].text = clean_val
                table_headers = []
                table_rows = []
                
        if not stripped:
            continue
            
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('#### '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[5:])
            run.bold = True
        elif stripped.startswith('- ') or stripped.startswith('* '):
            list_text = stripped[2:]
            clean_text = re.sub(r'\*\*(.*?)\*\*|_(.*?)_', r'\1\2', list_text)
            doc.add_paragraph(clean_text, style='List Bullet')
        elif re.match(r'^\d+\.\s', stripped):
            match = re.match(r'^(\d+)\.\s(.*)', stripped)
            list_text = match.group(2)
            clean_text = re.sub(r'\*\*(.*?)\*\*|_(.*?)_', r'\1\2', list_text)
            doc.add_paragraph(clean_text, style='List Number')
        else:
            clean_text = re.sub(r'\*\*(.*?)\*\*|_(.*?)_', r'\1\2', stripped)
            doc.add_paragraph(clean_text)
            
    doc.save(docx_path)
    print(f"Successfully compiled {md_path} to {docx_path}")

parse_markdown_to_docx(
    'Global_organization_subscription_functional_spec_session_first.md',
    'Global_organization_subscription_functional_spec_session_first.docx'
)

parse_markdown_to_docx(
    'Global_organization_subscription_walkthrough_session_first.md',
    'Global_organization_subscription_walkthrough_session_first.docx'
)

parse_markdown_to_docx(
    'Global_organization_subscription_db_linking_architecture.md',
    'Global_organization_subscription_db_linking_architecture.docx'
)
